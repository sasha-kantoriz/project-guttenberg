"""Guttenberg2.py.

usage: python3 guttenberg2.py [options]

Project Guttenberg books scrape script:

options:
  -h, --help            show this help message and exit
  -i INDEXES, --indexes INDEXES
                        books indexes to process, comma separated
  -s START, --start START
                        start index of the program
  -e END, --end END     end index of the program
  -w, --word            generate Word documents
  -c, --cover           generate PDF covers
  --interior            generate PDF interior only

Script will create output folder named as datestamp, and also maintain last processed book index and Excel file with each run spreadsheet
"""

import os
import re
import sys
import argparse
import pathlib
import requests
import fpdf
import docx
import openpyxl
from time import sleep
from datetime import datetime
from random import randint
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from pdf2image import convert_from_path
from bs4 import BeautifulSoup
from openai import OpenAI


client = OpenAI()


def search_open_library(title, author_name):
    base_url = 'http://openlibrary.org/search.json'
    params = {'title': title, 'author': author_name}

    try:
        response = requests.get(base_url, params=params)
        if response.status_code == 200:
            data = response.json()
            if 'docs' in data and len(data['docs']) > 0:
                book_data = data['docs'][0]

                # Extract publication year
                pub_year = book_data.get('first_publish_year', 'N/A')

                # Extract author key to get death year
                author_key = book_data['author_key'][0] if 'author_key' in book_data and book_data['author_key'] else None
                death_year = 'N/A'
                if author_key:
                    author_url = f"http://openlibrary.org/authors/{author_key}.json"
                    author_response = requests.get(author_url)
                    if author_response.status_code == 200:
                        author_data = author_response.json()
                        death_date = author_data.get('death_date', None)
                        if death_date:
                            try:
                                death_year = death_date.split('-')[0]
                                if not death_year.isdigit():
                                    death_year = 'N/A'
                            except:
                                death_year = 'N/A'

                return {'open_library_publication_year': pub_year, 'open_library_death_year': death_year}
            else:
                print(f"No book found for {title} by {author_name} on Open Library")
                return {'open_library_publication_year': 'N/A', 'open_library_death_year': 'N/A'}
        else:
            print(f"Failed to fetch Open Library data for {title} by {author_name}")
    except Exception as e:
        print(f"Error fetching data for {title} by {author_name} from Open Library: {e}")

    return {'open_library_publication_year': 'N/A', 'open_library_death_year': 'N/A'}


def search_wikipedia_author(author_name):
    try:
        search_url = "https://en.wikipedia.org/w/api.php"
        search_params = {'action': 'query', 'format': 'json', 'list': 'search', 'srsearch': author_name}
        response = requests.get(search_url, params=search_params)
        data = response.json()

        if 'query' in data and 'search' in data['query'] and data['query']['search']:
            page_title = data['query']['search'][0]['title']
            content_url = f"https://en.wikipedia.org/w/api.php"
            content_params = {"action": "parse", "format": "json", "page": page_title}
            response = requests.get(content_url, params=content_params)
            data = response.json()
            page_text = data['parse']['text']['*']

            death_year = 'N/A'
            death_match = re.search(r'\b(?:died|death)\s*(?:on|in)?\s*(\d{4})', page_text, re.IGNORECASE)
            if death_match:
                death_year = death_match.group(1)
                return death_year

    except Exception as e:
        print(f"Error fetching data for {author_name} from Wikipedia: {e}")

    return 'N/A'


def search_google_books(title, author_name, retries=3):
    base_url = 'https://www.googleapis.com/books/v1/volumes'
    params = {'q': f'intitle:{title}+inauthor:{author_name}'}

    for attempt in range(retries):
        try:
            response = requests.get(base_url, params=params)
            if response.status_code == 200:
                data = response.json()
                if 'items' in data and data['totalItems'] > 0:
                    book_data = data['items'][0]['volumeInfo']

                    pub_year = book_data.get('publishedDate', 'N/A')
                    if len(pub_year) > 4:
                        pub_year = pub_year.split('-')[0]

                    return {'google_books_publication_year': pub_year}
                else:
                    print(f"No book found for {title} by {author_name} on Google Books")
                    return {'google_books_publication_year': 'N/A'}
            else:
                print(f"Failed to fetch Google Books data for {title} by {author_name}, attempt {attempt+1}")
                sleep(1)
        except Exception as e:
            print(f"Error fetching data for {title} by {author_name} from Google Books: {e}")

    print(f"Failed to fetch data for {title} by {author_name} after {retries} attempts")
    return {'google_books_publication_year': 'N/A'}


def search_wikidata(author_name):
    base_url = 'https://www.wikidata.org/w/api.php'
    params = {'action': 'wbsearchentities', 'format': 'json', 'language': 'en', 'search': author_name, 'type': 'item'}

    try:
        response = requests.get(base_url, params=params)
        if response.status_code == 200:
            data = response.json()
            if 'search' in data and len(data['search']) > 0:
                author_id = data['search'][0]['id']
                author_url = f"https://www.wikidata.org/wiki/Special:EntityData/{author_id}.json"
                author_response = requests.get(author_url)
                if author_response.status_code == 200:
                    author_data = author_response.json()
                    entities = author_data.get('entities', {})
                    author_info = entities.get(author_id, {})
                    claims = author_info.get('claims', {})
                    death_date = claims.get('P570', [{}])[0].get('mainsnak', {}).get('datavalue', {}).get('value', {}).get('time', None)
                    if death_date:
                        death_year = death_date[1:5]
                        return death_year
            else:
                print(f"No data found on Wikidata for {author_name}")
        else:
            print(f"Failed to fetch Wikidata data for {author_name}")
    except Exception as e:
        print(f"Error fetching data for {author_name} from Wikidata: {e}")

    return 'N/A'


class PDF(fpdf.FPDF):
    def footer(self):
        if self.page_no() != 1:
            # Go to 1.5 cm from bottom
            self.set_y(-15)
            self.add_font("dejavu-sans", style="", fname="assets/DejaVuSans.ttf")
            self.set_font("dejavu-sans", size=8)
            # Print centered page number
            self.cell(0, 10, f"{self.page_no()}", 0, 0, 'C')


def get_latest_published_book_index():
    url = 'https://www.gutenberg.org/ebooks/search/?sort_order=release_date'
    response = requests.get(url, timeout=60)
    html = BeautifulSoup(response.content, features="html.parser")
    latest_book = html.body.find('li', attrs={'class': 'booklink'})
    index = latest_book.a.attrs['href'].split('/')[-1]
    return int(index)


def update_last_index(index):
    with open('index', 'w') as f:
        f.write(str(index))


def get_previous_last_index():
    try:
        return int(open('index').read()) + 1
    except:
        return 1


def generate_book_pdfs(folder, _id, title, author, description, notes, contents, preface, text, appendix, interior_only=False, cover_only=False, word_only=False):
    interior_pdf_fname, cover_pdf_fname, front_cover_pdf_fname, front_cover_webp_fname, front_cover_square_fname, front_cover_image_tmp_fname, dalle_cover_img_png, dalle_cover_img_webp = (
        f"{folder}/pdf/{_id}_paperback_interior.pdf",
        f"{folder}/cover/{_id}_paperback_cover.pdf",
        f"{folder}/front_cover//{_id}.pdf",
        f"{folder}/front_cover/{_id}.webp",
        f"{folder}/front_cover/{_id}_square.webp",
        f"{folder}/front_cover/{_id}.png",
        f"{folder}/imgs/{_id}.png",
        f"{folder}/imgs/{_id}.webp"
    )
    pdf = PDF(format=(152.4, 228.6))
    pdf.add_font("dejavu-sans", style="", fname="assets/DejaVuSans.ttf")
    # TITLE
    pdf.add_page()
    pdf.set_font("dejavu-sans", size=24)
    lines_num = len(pdf.multi_cell(w=0, align='C', padding=(0, 8), text=f"{title}\n\n{author}", dry_run=True, output="LINES"))
    if lines_num >= 3:
        padding_top = (228.6 - 24 * (lines_num - 1)) / 2
    else:
        padding_top = (228.6 - 24 * (lines_num)) / 2
    pdf.multi_cell(w=0, align='C', padding=(padding_top, 8, 0), text=f"{title}\n\n{author}")
    # PUBLISHER NOTES
    if notes:
        pdf.add_page()
        pdf.set_font("dejavu-sans", size=10)
        pdf.multi_cell(w=0, h=4, align='J', padding=8, text=notes)
    # CONTENTS
    if contents:
        pdf.add_page()
        pdf.set_font("dejavu-sans", size=10)
        pdf.multi_cell(w=0, h=4.4, align='J', padding=8, text=contents)
    # PREFACE
    if preface:
        pdf.add_page()
        pdf.set_font("dejavu-sans", size=10)
        pdf.multi_cell(w=0, h=4.4, align='J', padding=8, text=preface)
    # TEXT
    pdf.add_page()
    pdf.set_font("dejavu-sans", size=10)
    pdf.multi_cell(w=0, h=4.4, align='J', padding=8, text=text)
    # INDEX
    if appendix:
        pdf.add_page()
        pdf.set_font("dejavu-sans", size=10)
        pdf.multi_cell(w=0, h=4.4, align='J', padding=8, text=appendix)
    #
    pages = pdf.page_no()
    if 24 <= pages <= 828 and not cover_only and not word_only:
        pdf.output(interior_pdf_fname)
    # COVERS
    if 24 <= pages <= 828 and not (word_only or interior_only):
        # FRONT COVER
        cover_width, cover_height = 152.4 + 3.175, 234.95
        pdf = fpdf.FPDF(format=(cover_width, cover_height))
        pdf.add_font('dejavu-sans', style="", fname="assets/DejaVuSans.ttf")
        pdf.add_page()
        pdf.set_fill_color(r=250, g=249, b=222)
        pdf.rect(h=pdf.h, w=pdf.w, x=0, y=0, style="DF")
        pdf.set_font('dejavu-sans', size=18)
        text_h = pdf.multi_cell(w=0, align='C', padding=6.35, text=f"\n\n{title}\n* * *\n{author}\n", dry_run=True, output="HEIGHT")
        pdf.multi_cell(w=0, align='C', padding=6.35, text=f"\n\n{title}\n\n* * *\n\n{author}\n")
        # COVER IMAGE
        include_cover_img = (text_h + 8) < 234.95 - ((234.95 - 40) / 2 + 10)
        #
        if include_cover_img:
            try:
                prompt = f"Generate an image to be featured in a book cover. Exclude any depictions of books, book covers or written text. Meeting the criteria mentioned before, the image needs to be based on the following description: {description}"
                img_url = client.images.generate(model='dall-e-3', prompt=prompt, n=1, quality="standard").data[0].url
                response = requests.get(img_url)
                with open(dalle_cover_img_png, 'wb') as img:
                    img.write(response.content)
                pdf.image(dalle_cover_img_png, x=(152.4 - 100 + 6.35) / 2,
                          y=(234.95 - 40) / 2 + 20, w=100, h=100)
            except:
                pass
        pdf.output(front_cover_pdf_fname)
        try:
            front_cover_pages = convert_from_path(front_cover_pdf_fname)
            front_cover_pages[0].save(front_cover_image_tmp_fname, "PNG")
            cover_webp = Image.open(front_cover_image_tmp_fname)
            width, height = cover_webp.size
            max_dim = max(width, height)
            square_webp = Image.new('RGB', (max_dim, max_dim), (255, 255, 255))
            square_webp.paste(cover_webp, ((max_dim - width) // 2, (max_dim - height) // 2))
            square_webp.save(front_cover_square_fname, "WEBP")
            cover_webp.save(front_cover_webp_fname, "WEBP")
            cover_webp.close()
            cover_image = Image.open(dalle_cover_img_png)
            cover_image.save(dalle_cover_img_webp, "WEBP")
            cover_image.close()
        except:
            pass
        # Full cover
        cover_width, cover_height = 152.4 * 2 + pages * 0.05720 + 3.175 * 2, 234.95
        pdf = fpdf.FPDF(format=(cover_width, cover_height))
        pdf.add_font('dejavu-sans', style="", fname="assets/DejaVuSans.ttf")
        pdf.add_page()
        pdf.set_fill_color(r=250,g=249,b=222)
        pdf.rect(h=pdf.h, w=pdf.w, x=0, y=0, style="DF")
        cols = pdf.text_columns(ncols=2, gutter=pages*0.05720 + 1.588*2, l_margin=6.35, r_margin=6.35)
        #
        description_p = cols.paragraph(text_align='L')
        pdf.set_font('dejavu-sans', size=12)
        description_lines = pdf.multi_cell(w=152.4, align='L', padding=(0, 11.175), text=description, dry_run=True, output="LINES")
        description_p.write('\n'.join(description_lines))
        cols.end_paragraph()
        #
        cols.new_column()
        #
        title_p = cols.paragraph(text_align='C')
        pdf.set_font('dejavu-sans', size=24)
        title_h = pdf.multi_cell(w=0, align='C', padding=(0, 8), text=f"\n\n{title}", dry_run=True, output="HEIGHT")
        title_p.write(f"\n\n{title}")
        cols.end_paragraph()
        #
        separator_text = "\n* * *"
        separator_p = cols.paragraph(text_align='C')
        pdf.set_font('dejavu-sans', size=16)
        separator_h = pdf.multi_cell(w=0, align='C', padding=(0, 8), text=separator_text, dry_run=True, output="HEIGHT")
        separator_p.write(separator_text)
        cols.end_paragraph()
        #
        author_p = cols.paragraph(text_align='C')
        pdf.set_font('dejavu-sans', size=16)
        author_h = pdf.multi_cell(w=0, align='C', padding=(0, 8), text=f"\n{author}\n", dry_run=True, output="HEIGHT")
        author_p.write(f"\n{author}")
        cols.end_paragraph()
        #
        if include_cover_img:
            try:
                pdf.image(dalle_cover_img_png, x=(152.4 + pages * 0.05720 + 3.175) + (152.4 - 100 - 6.35) / 2 + 5, y=(234.95 - 40) / 2, w=100, h=100)
            except:
                pass
        #
        cols.render()
        pdf.output(cover_pdf_fname)
        try:
            os.remove(front_cover_pdf_fname)
            os.remove(front_cover_image_tmp_fname)
            os.remove(dalle_cover_img_png)
        except:
            pass
    #
    return (
        interior_pdf_fname,
        cover_pdf_fname,
        front_cover_webp_fname,
        pages
    )


def generate_book_docx(folder, _id, title, author, description, book_publisher_notes, preface, contents, text):
    doc = docx.Document("assets/template.docx")
    currentYear, currentMonth = datetime.now().year, datetime.now().month
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(f"{title}\n\n{author}")
    title_font = title_run.font
    title_font.name = 'Verdana'
    title_font.size = docx.shared.Pt(24)
    doc.add_page_break()
    if book_publisher_notes:
        preface_paragraph = doc.add_paragraph()
        preface_run = preface_paragraph.add_run(book_publisher_notes)
        preface_font = preface_run.font
        preface_font.name = 'Verdana'
        preface_font.size = docx.shared.Pt(10)
        doc.add_page_break()
    if preface:
        preface_paragraph = doc.add_paragraph()
        preface_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
        preface_run = preface_paragraph.add_run(preface)
        preface_font = preface_run.font
        preface_font.name = 'Verdana'
        preface_font.size = docx.shared.Pt(10)
        doc.add_page_break()
    if contents:
        contents_paragraph = doc.add_paragraph()
        contents_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
        contents_run = contents_paragraph.add_run(contents)
        contents_font = contents_run.font
        contents_font.name = 'Verdana'
        contents_font.size = docx.shared.Pt(10)
        doc.add_page_break()
    text_paragraph = doc.add_paragraph()
    text_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
    text_run = text_paragraph.add_run(text)
    text_font = text_run.font
    text_font.name = 'Verdana'
    text_font.size = docx.shared.Pt(10)
    doc.save(f"{folder}/word/{_id}_paperback_interior.docx")


def get_books(run_folder, start, end, interior_only=False, cover_only=False, word_only=False, indexes=None):
    update_index_flag = True
    datestamp = datetime.now().strftime('%Y-%B-%d %H_%M')
    if not (interior_only or cover_only or word_only):
        try:
            wb = openpyxl.load_workbook('Project Guttenberg.xlsx')
        except:
            wb = openpyxl.Workbook()
        try:
            del wb['Sheet']
        except:
            pass
        ws = wb.create_sheet(datestamp)
        ws.append(
            [
                "Book ID",
                "Plain text URL",
                "Title",
                # "Published Year",
                "Language",
                "Author",
                # "Author Year of Death",
                "Translator",
                "Illustrator",
                "Description",
                "Keywords",
                "BISAC codes",
                "Pages num",
                "PDF file name",
                "Cover PDF file name",
                "Front cover WEBP file name",
                "Google Book Publication Year",
                "OpenLibrary Book Publication Year",
                "Wikidata Author Year of Death",
                "Wikipedia Author Year of Death",
                "OpenLibrary Author Year of Death",
            ]
        )
    try:
        sequence = indexes if indexes else range(start, end + 1)
        for i in sequence:
            print(f'Processing index: {i}')
            book_url = f'https://www.gutenberg.org/ebooks/{i}.txt.utf-8'
            response = requests.get(
                book_url,
                timeout=60,
                headers={
                    'User-Agent': 'Mozilla/5.0 (Windows; U; Windows NT 6.1; zh-CN) AppleWebKit/533+ (KHTML, like Gecko)'
                }
            )
            #
            if response.status_code != 200:
                print(f"Error fetching book text: {response.status_code}")
                continue
            #
            book_txt = response.content.decode('utf-8')
            #
            book_author = re.search(r"Author: (.*)\n", book_txt)
            book_author = book_author.groups()[0] if book_author else ""
            book_author = book_author.strip().replace('\\', '-').replace('/', '-').replace('&', ' and ')
            book_language = re.search(r"Language: (.*)\n", book_txt, re.IGNORECASE)
            book_language = book_language.groups()[0] if book_language else ""
            book_translator = re.search(r"Translator: (.*)\n", book_txt, re.IGNORECASE)
            book_translator = book_translator.groups()[0] if book_translator else ""
            book_illustrator = re.search(r"Illustrator: (.*)\n", book_txt, re.IGNORECASE)
            book_illustrator = book_illustrator.groups()[0] if book_illustrator else ""
            book_title = re.search(r"Title: (.*)\n", book_txt)
            book_title = book_title.groups()[0] if book_title else ""
            book_title = book_title.strip().replace('\\', '-').replace('/', '-').replace('&', ' and ')
            book_content_start_index = re.search(r"\*\*\* START OF THE PROJECT GUTENBERG .* \*\*\*", book_txt, re.IGNORECASE)
            book_content_start_index = book_content_start_index.end() if book_content_start_index else 0
            book_content_end_index = re.search(r"\*\*\* END OF THE PROJECT GUTENBERG .* \*\*\*", book_txt, re.IGNORECASE)
            book_content_end_index = book_content_end_index.start() if book_content_end_index else -1
            book_txt = book_txt[book_content_start_index:book_content_end_index]
            #
            if "hungarian" in book_language.lower() or "romanian" in book_language.lower() or "esperanto" in book_language.lower() or "latin" in book_language.lower() or "greek" in book_language.lower() or "tagalog" in book_language.lower() or "japanese" in book_language.lower() or "slovenian" in book_language.lower() or "telugu" in book_language.lower() or "Gaelic, Scottish" in book_language.lower() or "French, Dutch" in book_language.lower() or "English, Spanish" in book_language.lower() or "ojibwa" in book_language.lower() or not book_author or book_translator or book_illustrator:
                continue
            #
            illustrations_patterns = [
                re.compile(r'\[(\s+)?Illustration.+?](\r\n){2}', re.IGNORECASE|re.DOTALL),
                re.compile(r'\[(\s+)?Ilustracion.+?](\r\n){2}', re.IGNORECASE|re.DOTALL),
                re.compile(r'\[(\s+)?Ilustración.+?](\r\n){2}', re.IGNORECASE|re.DOTALL),
            ]
            for _pattern in illustrations_patterns:
                book_txt = re.sub(_pattern, '', book_txt)
            proofread_patterns = [
                re.compile(r'Produced(.+?)?(\s+)?at(\s+)?(https://|http://)?(www\.)?pgdp\.net(\s+)?(.+?)?(\r\n){4}', re.IGNORECASE|re.DOTALL),
                re.compile(r'Produced(.+?)?(\s+)?by(\s+)?(www\.)?ebooksgratuits\.com(\s+)?(.+?)?(\r\n){4}', re.IGNORECASE|re.DOTALL),
                re.compile(r'E(-)?(text|book)(\s+)?(produced|prepared)(\s+)?(.+?)?(\r\n){4}', re.IGNORECASE|re.DOTALL),
            ]
            for _pattern in proofread_patterns:
                book_txt = re.sub(_pattern, '', book_txt)
            transcriber_notes_patterns = [
                re.compile(r'(\[)?Transcriber(.+?)?(\s+)?Note(s)?(\s+)?(:)?(\s+)?(.+?)?(\r\n){4}', re.IGNORECASE | re.DOTALL),
                re.compile(r'\[Sidenote(s)?(\s+)?:(\s+)?(.+?)?(\r\n){2}', re.IGNORECASE | re.DOTALL),
                re.compile(r'\[Note(s)?(\s+)?:(\s+)?(.+?)?(\r\n){2}', re.IGNORECASE | re.DOTALL),
            ]
            for _pattern in transcriber_notes_patterns:
                book_txt = re.sub(_pattern, '', book_txt)
            start_end_patterns = [
                re.compile(r'START(\s+)?OF(\s+)?(THE)?(\s+)?PROJECT(\s+)?GUTENBERG.+?(\r\n){4}', re.IGNORECASE|re.DOTALL),
                re.compile(r'END(\s+)?OF(\s+)?(THE)?(\s+)?PROJECT(\s+)?GUTENBERG.+?(\r\n){4}', re.IGNORECASE|re.DOTALL),
            ]
            for _pattern in start_end_patterns:
                book_txt = re.sub(_pattern, '', book_txt)
            #
            book_txt = book_txt.replace('\r\n', '\n')
            # BOOK PUBLISHER NOTES
            book_publisher_notes_start_index, book_publisher_notes_end_index = 0, book_txt[100:].find('\n\n\n\n')
            if book_publisher_notes_end_index != -1:
                book_publisher_notes_end_index += 100
            else:
                book_publisher_notes_end_index = 0
            book_publisher_notes = book_txt[book_publisher_notes_start_index:book_publisher_notes_end_index]
            # BOOK CONTENTS
            contents_search = re.search(r"(content|contents|chapters|file numbers)(:)?(\.)?(\n){2}", book_txt, re.IGNORECASE)
            if contents_search and not re.search(r"(content|contents|chapters|file numbers)(:)?(\.)?(\n)+(\s)*of", book_txt[:contents_search.start() + 100], re.IGNORECASE):
                contents_start_index = contents_search.start()
                contents_end_index = contents_start_index + len(contents_search.group()) + 5 + book_txt[contents_start_index + len(contents_search.group()) + 5:].find('\n\n\n\n')
            else:
                contents_end_index = contents_start_index = 0
            book_contents = book_txt[contents_start_index:contents_end_index]
            preface_search = re.search(r'(preface|foreword)(\.)?(\n){2}', book_txt, re.IGNORECASE)
            if preface_search:
                preface_start_index = preface_search.start()
                preface_end_index = preface_start_index + len(preface_search.group()) + 10 + book_txt[preface_start_index + len(preface_search.group()) + 10:].find('\n\n\n\n')
                book_preface = book_txt[preface_start_index:preface_end_index]
            else:
                preface_end_index = 0
                book_preface = ""
            # BOOK INDEX
            appendix_search = re.search(r'(INDEX|Index|APPENDIX|Appendix)(\.)?(:)?(\n){2}', book_txt)
            if appendix_search:
                appendix_start_index = appendix_search.start()
                appendix_end_index = appendix_start_index + len(appendix_search.group()) + 10 + book_txt[appendix_start_index + len(appendix_search.group()) + 10:].find('\n\n\n\n')
                book_appendix = book_txt[appendix_start_index:appendix_end_index]
            else:
                appendix_start_index = len(book_txt)
                book_appendix = ""
            #
            book_txt = book_txt[max(contents_end_index, preface_end_index):appendix_start_index]
            #
            if book_txt.find('LIST OF ILLUSTRATIONS') != -1:
                illustrations_start_index = book_txt.find('LIST OF ILLUSTRATIONS')
                illustrations_end_index = illustrations_start_index + book_txt[illustrations_start_index:].find('\n\n\n\n')
                book_txt = book_txt[illustrations_end_index:]
            if book_contents in book_publisher_notes:
                book_publisher_notes = ""
            book_publisher_notes = book_publisher_notes.replace('\n\n\n\n', '\n\n').replace('_', '').replace('  ', ' ').replace('--', '-').replace('\n\n', '_____').replace('\n', ' ').replace('_____', '\n\n')
            book_contents = book_contents.replace('\n\n\n', '\n').replace('\n\n', '\n').replace('_', '').replace('  ', ' ').replace('--', '-')
            book_preface = book_preface.replace('\n\n\n\n', '\n\n').replace('_', '').replace('  ', ' ').replace('--', '-').replace('\n\n', '_____').replace('\n', ' ').replace('_____', '\n\n')
            book_txt = book_txt.replace('\n\n\n\n', '\n\n').replace('_', '').replace('  ', ' ').replace('--', '-').replace('\n\n', '_____').replace('\n', ' ').replace('_____', '\n\n')
            book_appendix = book_appendix.replace('\n\n\n\n', '\n\n').replace('_', '').replace('  ', ' ').replace('--', '-').replace('\n\n', '_____').replace('\n', ' ').replace('_____', '\n\n')
            ############################################################################################################
            # Book Metadata
            ############################################################################################################
            description_query = f"Provide a 150 words description of the classic book {book_title}"
            if book_author:
                description_query += f" by Author and Writer {book_author}."
            if book_language:
                description_query += f" Write the review in this language: {book_language}"
            description_completion = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": description_query
                    },
                ]
            )
            description = description_completion.choices[0].message.content

            ############################################################################################################
            # Book files Generation
            ############################################################################################################
            try:
                book_fname, cover_fname, front_cover_image_fname, pages_num = generate_book_pdfs(
                    run_folder, i, book_title, book_author, description, book_publisher_notes, book_contents, book_preface, book_txt, book_appendix, interior_only,cover_only, word_only
                )
                if (24 <= pages_num <= 828) and (not (cover_only or interior_only) or word_only):
                    generate_book_docx(
                        run_folder, i, book_title, book_author, description, book_publisher_notes, book_contents, book_preface, book_txt
                    )
                #
                if not (interior_only or cover_only or word_only):
                    keywords_query = f'Give me 7 keywords separated by semicolons (only the keywords, no numbers nor introductory words) that accurately reflect the main themes and genre of the classic book "{book_title}" by Author "{book_author}". Keywords must not be subjective claims about its quality, time-sensitive statments and must not include the word "book". Keywords must also not contain words included on the book the title, author nor contained on the following book description: {description}'
                    keywords_completion = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {
                                "role": "system",
                                "content": keywords_query
                            },
                        ]
                    )
                    keywords = keywords_completion.choices[0].message.content
                    #
                    bisac_codes_query = f'Give me up to 3 BISAC codes separated by semicolons (only the code in the official format, not its description and not numbered) for the book "{book_title}" by Author "{book_author}" with description "{description}", for its correct classification. Output format example would be: FIC019000; FIC031010; FIC014000'
                    bisac_codes_completion = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {
                                "role": "system",
                                "content": bisac_codes_query
                            },
                        ]
                    )
                    bisac_codes = bisac_codes_completion.choices[0].message.content
                    #
                    """
                    published_year_query = f'Please, tell me the year the book {book_title} by {book_author} was published. Provide only the date in the format YYYY.'
                    published_year_completion = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {
                                "role": "system",
                                "content": published_year_query
                            },
                        ]
                    )
                    published_year = published_year_completion.choices[0].message.content
                    #
                    author_year_of_death_query = f'Please, tell me the year of death of {book_author}, the author of the book {book_title}. Provide only the date in the format YYYY. If the author is still alive, please, provide the "XXXX".'
                    author_year_of_death_completion = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {
                                "role": "system",
                                "content": author_year_of_death_query
                            },
                        ]
                    )
                    author_year_of_death = author_year_of_death_completion.choices[0].message.content
                    """
                    # Extended Metadata
                    google_books_search_data = search_google_books(book_title, book_author)
                    open_library_search_data = search_open_library(book_title, book_author)
                    wikipedia_author_year_of_death = search_wikipedia_author(book_author)
                    wikidata_author_year_of_death = search_wikidata(book_author)
                    #
                    ws.append(
                        [
                            i,
                            book_url,
                            book_title,
                            # published_year,
                            book_language,
                            book_author,
                            # author_year_of_death,
                            book_translator,
                            book_illustrator,
                            description,
                            keywords,
                            bisac_codes,
                            pages_num,
                            book_fname,
                            cover_fname,
                            front_cover_image_fname,
                            google_books_search_data.get('google_books_publication_year', 'N / A'),
                            open_library_search_data.get('open_library_publication_year', 'N / A'),
                            wikidata_author_year_of_death,
                            wikipedia_author_year_of_death,
                            open_library_search_data.get('open_library_death_year', 'N / A'),
                        ]
                    )
            except:
                import traceback
                print(traceback.format_exc())
    except KeyboardInterrupt:
        update_index_flag = False
    except Exception as e:
        print(e)
        update_index_flag = False
        update_last_index(i)
    finally:
        if not (interior_only or word_only or cover_only):
            wb.save('Project Guttenberg.xlsx')
        # update last published book index
        if update_index_flag:
            update_last_index(end)


def parse_args():
    # parse command line arguments
    parser = argparse.ArgumentParser(
        prog='guttenberg2.py',
        usage='python3 %(prog)s [options]',
        description='Project Guttenberg books scrape script:',
        epilog="Script will create output folder named as datestamp, and also maintain last processed book index and Excel file with each run spreadsheet"
    )
    parser.add_argument('-i', '--indexes', type=str, dest='indexes', default='', help='books indexes to process, comma separated')
    parser.add_argument('-s', '--start', type=int, dest='start', default=get_previous_last_index(),
                        help='start index of the program')
    parser.add_argument('-e', '--end', type=int, dest='end', default=get_latest_published_book_index(),
                        help='end index of the program')
    parser.add_argument('-w', '--word', action='store_true', help='generate Word documents')
    parser.add_argument('-c', '--cover', action='store_true', help='generate PDF covers')
    parser.add_argument('--interior', action='store_true', help='generate PDF interior only')
    #
    return parser.parse_args()


if __name__ == '__main__':
    # create PDFs output folder
    run_folder = datetime.now().strftime('%Y-%B')
    pathlib.Path(f"{run_folder}/imgs").mkdir(parents=True, exist_ok=True)
    pathlib.Path(f"{run_folder}/cover").mkdir(parents=True, exist_ok=True)
    pathlib.Path(f"{run_folder}/front_cover").mkdir(parents=True, exist_ok=True)
    pathlib.Path(f"{run_folder}/word").mkdir(parents=True, exist_ok=True)
    pathlib.Path(f"{run_folder}/pdf").mkdir(parents=True, exist_ok=True)
    #
    args = parse_args()
    get_books(run_folder, args.start, args.end, args.interior, args.cover, args.word, args.indexes.split(',') if args.indexes else None)
